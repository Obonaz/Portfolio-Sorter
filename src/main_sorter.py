import os
import logging

from .logger_setup import setup_logging, LOGGER_NAME
from .file_utils import get_source_files, create_target_subdir, move_file_to_directory
from .text_extractor import extract_text, TextExtractionError, ALL_SUPPORTED_EXTENSIONS
from .categorizer import LLMCategorizer, PREDEFINED_CATEGORIES

# Configure logging for the application
# This will set up the handlers and formatters.
# If main_sorter is imported elsewhere, this ensures logger is ready.
setup_logging() 
logger = logging.getLogger(LOGGER_NAME)

EXCLUDED_FILE_EXTENSIONS = [
    '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.svg',  # Images
    '.zip', '.rar', '.7z', '.tar', '.gz',  # Archives
    '.mp3', '.wav', '.aac', '.flac',  # Audio
    '.mp4', '.avi', '.mov', '.mkv',  # Video
    '.exe', '.dll', '.so', '.dmg', '.app', # Executables & libraries
    '.iso', # Disk images
    # Add any other non-document extensions you want to explicitly exclude from processing
]

def process_files(source_dir: str, target_dir: str) -> None:
    """
    Orchestrates the file sorting process:
    1. Scans source directory for files.
    2. Extracts text from supported document types.
    3. Categorizes extracted text.
    4. Moves categorized files to appropriate subdirectories in the target directory.
    """
    categorizer = LLMCategorizer()

    logger.info(f"Starting file sorting process. Source: '{source_dir}', Target: '{target_dir}'")

    if not os.path.isdir(source_dir):
        logger.error(f"Source directory '{source_dir}' does not exist. Aborting process.")
        return
    if not os.path.isdir(target_dir):
        logger.warning(f"Target directory '{target_dir}' does not exist. It will be created.")
        try:
            os.makedirs(target_dir, exist_ok=True)
            logger.info(f"Target directory '{target_dir}' created.")
        except OSError as e:
            logger.error(f"Failed to create target directory '{target_dir}': {e}. Aborting process.")
            return

    try:
        files_to_process = get_source_files(source_dir, EXCLUDED_FILE_EXTENSIONS)
        logger.info(f"Found {len(files_to_process)} files to process (after exclusions).")
    except FileNotFoundError: # Already handled by the check above, but good practice for get_source_files
        logger.error(f"Source directory '{source_dir}' not found during file scan. Aborting.")
        return
    except Exception as e:
        logger.error(f"An unexpected error occurred while getting source files: {e}. Aborting.")
        return

    for file_path in files_to_process:
        logger.info(f"Processing file: {file_path}")
        file_name = os.path.basename(file_path)
        file_ext = os.path.splitext(file_name)[1].lower()

        # Text Extraction
        try:
            if file_ext not in ALL_SUPPORTED_EXTENSIONS:
                logger.info(f"File '{file_name}' (type: {file_ext}) is not a supported document type for text extraction. Skipping categorization, leaving in source.")
                continue

            text_content = extract_text(file_path)
            if text_content is None or not text_content.strip():
                logger.warning(f"Text extraction returned None or empty text for supported file '{file_name}'. Skipping.")
                continue
            # Optional: logger.debug(f"Extracted text from '{file_name}': {text_content[:200]}...")
        
        except TextExtractionError as e:
            logger.error(f"Failed to extract text from '{file_name}': {e}. Leaving in source folder.")
            continue
        except FileNotFoundError: # Should be caught by get_source_files, but as a safeguard for extract_text
            logger.error(f"File '{file_name}' not found during text extraction. This should not happen if previously listed. Leaving in source folder.")
            continue
        except Exception as e:
            logger.error(f"An unexpected error occurred during text extraction for '{file_name}': {e}. Leaving in source folder.")
            continue

        # Categorization
        category = categorizer.categorize_text(text_content, PREDEFINED_CATEGORIES)

        # File Organization
        if category:
            try:
                target_category_dir = create_target_subdir(target_dir, category)
                move_file_to_directory(file_path, target_category_dir)
                logger.info(f"Moved '{file_name}' to '{target_category_dir}'. Category: {category}")
            except FileNotFoundError: # If source file somehow disappeared
                 logger.error(f"File '{file_name}' not found during move operation. File may have been moved or deleted externally.")
            except OSError as e: # For issues with create_target_subdir or move_file_to_directory related to OS operations
                logger.error(f"OS error while creating subdirectory or moving '{file_name}' to category '{category}': {e}. File remains in source.")
            except Exception as e:
                logger.error(f"Failed to move '{file_name}' to category '{category}': {e}. File remains in source.")
        else:
            logger.info(f"File '{file_name}' could not be categorized. Leaving in source folder.")

    logger.info("File sorting process completed.")

if __name__ == '__main__':
    # Example Usage (for testing purposes)
    # Create dummy directories and files for testing
    
    TEST_SOURCE_DIR = "temp_source_main_sorter"
    TEST_TARGET_DIR = "temp_target_main_sorter"
    
    # Clean up previous test runs
    if os.path.exists(TEST_SOURCE_DIR):
        import shutil
        shutil.rmtree(TEST_SOURCE_DIR)
    if os.path.exists(TEST_TARGET_DIR):
        import shutil
        shutil.rmtree(TEST_TARGET_DIR)
    
    os.makedirs(TEST_SOURCE_DIR, exist_ok=True)
    os.makedirs(os.path.join(TEST_SOURCE_DIR, "subfolder"), exist_ok=True) # Test subdirectories
    os.makedirs(TEST_TARGET_DIR, exist_ok=True)

    # Create some dummy files
    dummy_files_content = {
        "report_final.docx": "This is a final report about project success.",
        "old_assignment.pdf": "Coursework assignment details here.",
        "meeting_slides.pptx": "Presentation slides for the weekly meeting.",
        "important_notes.txt": "These are some important notes, but .txt is not in ALL_SUPPORTED_EXTENSIONS by default.",
        "archive.zip": "some binary data",
        "image.jpg": "some image data",
        "thesis_document.docx": "My full thesis on advanced sorting algorithms.",
        "empty.docx": "", # File with supported extension but no text after extraction (mocked)
        "unsupported.xyz": "Unknown file type content",
        "subfolder/quiz1_answers.pdf": "Answer key for the first quiz."
    }

    # For actual text extraction to work, these would need to be real files.
    # We are mocking the content for categorization, but extract_text will try to open them.
    # So, we need to ensure that extract_text can handle these.
    # The current extract_text will fail for non-valid office files.
    # For this test, we'll just create empty files for supported types,
    # and rely on the logic within process_files to handle TextExtractionError or empty content.

    logger.info(f"Setting up test files in {TEST_SOURCE_DIR}")
    for name, _ in dummy_files_content.items():
        full_path = os.path.join(TEST_SOURCE_DIR, name)
        os.makedirs(os.path.dirname(full_path), exist_ok=True) # Ensure subfolder exists
        with open(full_path, 'w') as f:
            # For DOCX, PDF, PPTX, create minimal valid files if possible, or just empty.
            # For this test, simple empty files will cause TextExtractionError or return None/empty.
             if name.endswith(tuple(ALL_SUPPORTED_EXTENSIONS)):
                f.write("") # Let extract_text try and fail or return empty.
             else:
                f.write("dummy content for non-supported types")
        logger.info(f"Created test file: {full_path}")
    
    # Create a valid docx file for testing at least one successful extraction and categorization
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("This is a test report document for categorization.")
        doc.save(os.path.join(TEST_SOURCE_DIR, "real_report.docx"))
        logger.info(f"Created real_report.docx for testing.")
        doc2 = Document()
        doc2.add_paragraph("This is my dissertation about advanced topics.")
        doc2.save(os.path.join(TEST_SOURCE_DIR, "my_dissertation.docx"))
        logger.info(f"Created my_dissertation.docx for testing.")

    except ImportError:
        logger.warning("python-docx not installed. Cannot create real .docx for testing. Some tests might not be accurate.")


    logger.info("Starting main_sorter test process...")
    process_files(TEST_SOURCE_DIR, TEST_TARGET_DIR)
    logger.info("Main_sorter test process finished.")

    # Add some assertions here to check if files were moved as expected
    # Example:
    expected_report_path = os.path.join(TEST_TARGET_DIR, "Reports", "real_report.docx")
    if os.path.exists(expected_report_path):
        logger.info(f"Test PASSED: 'real_report.docx' was correctly moved to Reports.")
    else:
        logger.error(f"Test FAILED: 'real_report.docx' was NOT moved to Reports. Check logs at {os.path.join(os.getcwd(), 'logs', 'sorting_app.log')}")

    expected_thesis_path = os.path.join(TEST_TARGET_DIR, "Theses & Dissertations", "my_dissertation.docx")
    if os.path.exists(expected_thesis_path):
        logger.info(f"Test PASSED: 'my_dissertation.docx' was correctly moved to Theses & Dissertations.")
    else:
        logger.error(f"Test FAILED: 'my_dissertation.docx' was NOT moved to Theses & Dissertations.")

    # Check that excluded files are still in source
    excluded_file_path = os.path.join(TEST_SOURCE_DIR, "archive.zip")
    if os.path.exists(excluded_file_path):
        logger.info(f"Test PASSED: Excluded file 'archive.zip' remains in source.")
    else:
        logger.error(f"Test FAILED: Excluded file 'archive.zip' was moved or deleted from source.")
    
    unsupported_file_path = os.path.join(TEST_SOURCE_DIR, "important_notes.txt")
    if os.path.exists(unsupported_file_path):
         logger.info(f"Test PASSED: File 'important_notes.txt' (not in ALL_SUPPORTED_EXTENSIONS) remains in source.")
    else:
        logger.error(f"Test FAILED: File 'important_notes.txt' was moved or deleted from source.")


    # Clean up test directories after a delay or manual inspection
    # print(f"Test files are in {TEST_SOURCE_DIR} and {TEST_TARGET_DIR}. Inspect and then press Enter to clean up.")
    # input()
    # shutil.rmtree(TEST_SOURCE_DIR)
    # shutil.rmtree(TEST_TARGET_DIR)
    # logger.info("Cleaned up test directories.")
